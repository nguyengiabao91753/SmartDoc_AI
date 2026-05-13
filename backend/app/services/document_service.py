from typing import Any, Dict, List
import os
import pdfplumber
import re
import docx
import json
import base64
from pathlib import Path
from PIL import Image
import io
from langchain_core.documents import Document

from app.core.config import settings
from app.core.logger import LOG
from underthesea import sent_tokenize
import tiktoken


try:
    import pytesseract
    # Configure Tesseract path if specified in settings
    if hasattr(settings, 'TESSERACT_CMD') and settings.TESSERACT_CMD:
        pytesseract.pytesseract_cmd = settings.TESSERACT_CMD
    OCR_AVAILABLE = True
except:
    OCR_AVAILABLE = False

enc = tiktoken.get_encoding("cl100k_base")

# Image directories
IMAGES_DIR = os.path.join(settings.DATA_DIR, "images")
os.makedirs(IMAGES_DIR, exist_ok=True)

def token_len(text: str) -> int:
    return len(enc.encode(text))

def trim_to_token_limit(text: str, max_tokens: int) -> str:
    tokens = enc.encode(text)
    return enc.decode(tokens[:max_tokens])

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def save_image_from_page(page, page_num: int, document_id: int, chunk_id: int) -> str:
    """
    Hàm này lưu ảnh của một trang PDF vào thư mục images/doc_{document_id}/chunk_{chunk_id}_page_{page_num}.png
    
    """
    try:
        img = page.to_image(resolution=200).original
        doc_img_dir = os.path.join(IMAGES_DIR, f"doc_{document_id}")
        os.makedirs(doc_img_dir, exist_ok=True)
        
        img_name = f"chunk_{chunk_id}_page_{page_num}.png"
        img_path = os.path.join(doc_img_dir, img_name)
        img.save(img_path, "PNG", quality=85)
        
        # Return relative path for database storage
        rel_path = os.path.join("images", f"doc_{document_id}", img_name)
        LOG.debug(f"Saved image: {rel_path}")
        return rel_path
    except Exception as e:
        LOG.warning(f"Failed to save image from page {page_num}: {e}")
        return None

def _ocr_embedded_images(page, page_num: int) -> str:
    """
    Dùng page.images của pdfplumber để crop từng ảnh nhúng → OCR → trả về text ghép.
    Bỏ qua ảnh quá nhỏ (icon, logo, decoration).
    """
    if not OCR_AVAILABLE:
        return ""

    embedded = page.images
    if not embedded:
        return ""

    LOG.debug(f"Page {page_num}: found {len(embedded)} embedded image(s)")
    ocr_parts = []

    for idx, img_info in enumerate(embedded):
        try:
            x0 = img_info.get("x0", 0)
            top = img_info.get("top", 0)
            x1 = img_info.get("x1", 0)
            bottom = img_info.get("bottom", 0)

            # Bỏ qua ảnh quá nhỏ (< 80px mỗi chiều)
            width_pt  = x1 - x0
            height_pt = bottom - top
            if width_pt < 80 or height_pt < 80:
                LOG.debug(f"  Image {idx}: too small ({width_pt:.0f}x{height_pt:.0f}pt), skipping")
                continue

            cropped   = page.crop((x0, top, x1, bottom))
            pil_img   = cropped.to_image(resolution=300).original

            img_text  = pytesseract.image_to_string(pil_img, lang="vie+eng", config="--psm 6")
            img_text  = clean_text(img_text)

            if img_text.strip():
                LOG.debug(f"  Image {idx}: OCR extracted {len(img_text)} chars")
                ocr_parts.append(img_text)

        except Exception as e:
            LOG.warning(f"  OCR failed on embedded image {idx}, page {page_num}: {e}")

    return "\n".join(ocr_parts)


def extract_text_from_page(
    page,
    page_num: int = None,
    document_id: int = None,
    chunk_id: int = None,
) -> Dict[str, Any]:
    """
    Extract text từ page theo thứ tự ưu tiên:
    1. Native text (pdfplumber)
    2. OCR ảnh nhúng trong trang  ← luôn chạy, không phụ thuộc bước 1
    3. OCR toàn trang              ← fallback khi cả 2 trên đều rỗng
    """
    native_text = clean_text(page.extract_text() or "")
    has_ocr     = False
    image_path  = None

    # Lưu ảnh trang để làm citation
    if document_id is not None and page_num is not None and chunk_id is not None:
        try:
            image_path = save_image_from_page(page, page_num, document_id, chunk_id)
        except Exception as e:
            LOG.warning(f"Failed to save image from page {page_num}: {e}")

    # ── Bước 2: OCR ảnh nhúng (luôn chạy dù có native text hay không) ────────
    embedded_ocr_text = _ocr_embedded_images(page, page_num)
    if embedded_ocr_text.strip():
        has_ocr = True

    # ── Bước 3: Fallback OCR toàn trang nếu cả native lẫn embedded đều rỗng ──
    full_page_ocr_text = ""
    if not native_text and not embedded_ocr_text and OCR_AVAILABLE:
        try:
            LOG.debug(f"Page {page_num}: no text found, running full-page OCR...")
            pil_img           = page.to_image(resolution=300).original
            full_page_ocr_text = clean_text(
                pytesseract.image_to_string(pil_img, lang="vie+eng", config="--psm 6")
            )
            if full_page_ocr_text.strip():
                has_ocr = True
                LOG.debug(f"Page {page_num}: full-page OCR extracted {len(full_page_ocr_text)} chars")
        except Exception as e:
            LOG.warning(f"Full-page OCR failed for page {page_num}: {e}")

    # ── Ghép text theo thứ tự: native → embedded OCR → full-page OCR ──────────
    final_parts = [t for t in [native_text, embedded_ocr_text, full_page_ocr_text] if t.strip()]
    final_text  = "\n\n".join(final_parts)

    return {
        "text":       final_text,
        "has_ocr":    has_ocr,
        "image_path": image_path,
    }

def process_pdf(file_path: str, document_id: int = None) -> List[Dict]:
    LOG.info(f"Processing PDF: {file_path}")
    chunks: List[Dict] = []
    chunk_id = 0
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
            LOG.info(f"Total pages: {num_pages}")
            current_sentences: List[str] = []
            current_token_count = 0
            current_page_start = 1
            chunk_images: List[str] = []
            chunk_has_ocr = False
            pages_with_no_text = 0
            all_ocr_texts: List[str] = []  # Collect ALL OCR text as fallback
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_extract = extract_text_from_page(page, page_num, document_id, chunk_id)
                page_text = page_extract["text"]
                page_has_ocr = page_extract["has_ocr"]
                page_image_path = page_extract["image_path"]
                
                # Track pages without native text
                if not page_text:
                    pages_with_no_text += 1
                    # Still track image path even if no text extracted
                    if page_image_path:
                        if not chunk_has_ocr:
                            chunk_has_ocr = True
                        chunk_images.append(page_image_path)
                    continue
                
                # Track if any part of chunk used OCR
                if page_has_ocr:
                    chunk_has_ocr = True
                    if page_image_path:
                        chunk_images.append(page_image_path)
                    # Collect OCR text for fallback
                    all_ocr_texts.append(page_text)
                
                sentences = sent_tokenize(page_text)
                for sentence in sentences:
                    sentence_tokens = token_len(sentence)
                    if sentence_tokens > settings.CHUNK_SIZE:
                        LOG.warning("Long sentence detected → trimming by token")
                        sentence = trim_to_token_limit(sentence, settings.CHUNK_SIZE)
                        sentence_tokens = token_len(sentence)
                    
                    if current_token_count + sentence_tokens > settings.CHUNK_SIZE:
                        if current_sentences:
                            chunk_text = " ".join(current_sentences)
                            chunks.append({
                                "id": chunk_id,
                                "text": chunk_text,
                                "tokens": current_token_count,
                                "page_start": current_page_start,
                                "page_end": page_num,
                                "has_ocr": chunk_has_ocr,
                                "image_paths": chunk_images.copy()
                            })
                            chunk_id += 1
                        
                        overlap_sentences = current_sentences[-settings.OVERLAP_SENTENCES:]
                        current_sentences = overlap_sentences.copy()
                        current_token_count = sum(token_len(s) for s in current_sentences)
                        current_page_start = page_num
                        chunk_images = []
                        chunk_has_ocr = False
                    
                    current_sentences.append(sentence)
                    current_token_count += sentence_tokens
            
            if current_sentences:
                chunk_text = " ".join(current_sentences)
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "tokens": current_token_count,
                    "page_start": current_page_start,
                    "page_end": num_pages,
                    "has_ocr": chunk_has_ocr,
                    "image_paths": chunk_images
                })
            
            # If no chunks but have OCR texts collected, create chunk from all OCR text
            if not chunks and all_ocr_texts:
                LOG.warning(f"Creating chunk from collected OCR text ({len(all_ocr_texts)} pages)")
                combined_ocr_text = " ".join(all_ocr_texts)
                chunks.append({
                    "id": 0,
                    "text": combined_ocr_text,
                    "tokens": sum(token_len(t) for t in all_ocr_texts),
                    "page_start": 1,
                    "page_end": num_pages,
                    "has_ocr": True,
                    "image_paths": chunk_images
                })
            # Fallback: if still no chunks but have images, create placeholder
            elif not chunks and chunk_images:
                LOG.warning(f"PDF has {pages_with_no_text} pages but no extractable text. Creating image-only chunk.")
                chunks.append({
                    "id": 0,
                    "text": f"[Scanned document - {num_pages} page(s) with images for visual reference]",
                    "tokens": 10,
                    "page_start": 1,
                    "page_end": num_pages,
                    "has_ocr": True,
                    "image_paths": chunk_images
                })
    except Exception as e:
        LOG.error(f"PDF processing error: {e}")
        return []
    
    LOG.info(f"Generated {len(chunks)} chunks, {pages_with_no_text} pages without text")
    return chunks



def process_docx(file_path: str, document_id: int = None) -> List[Dict]:
    LOG.info(f"Processing DOCX: {file_path}")
    doc = docx.Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    sentences = sent_tokenize(full_text)
    chunks: List[Dict] = []
    chunk_id = 0
    current_sentences: List[str] = []
    current_token_count = 0
    
    for sentence in sentences:
        sentence_tokens = token_len(sentence)
        if sentence_tokens > settings.CHUNK_SIZE:
            sentence = trim_to_token_limit(sentence, settings.CHUNK_SIZE)
            sentence_tokens = token_len(sentence)
        if current_token_count + sentence_tokens > settings.CHUNK_SIZE:
            if current_sentences:
                chunk_text = " ".join(current_sentences)
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "tokens": current_token_count,
                    "has_ocr": False,
                    "image_paths": []
                })
                chunk_id += 1
            overlap_sentences = current_sentences[-settings.OVERLAP_SENTENCES:]
            current_sentences = overlap_sentences.copy()
            current_token_count = sum(token_len(s) for s in current_sentences)
        current_sentences.append(sentence)
        current_token_count += sentence_tokens
    
    if current_sentences:
        chunk_text = " ".join(current_sentences)
        chunks.append({
            "id": chunk_id,
            "text": chunk_text,
            "tokens": current_token_count,
            "has_ocr": False,
            "image_paths": []
        })
        
    LOG.info(f"Generated {len(chunks)} chunks from DOCX")
    return chunks


class DocumentService:
    """
    Lớp tương thích để ráp nối logic xử lý file với RAGService.
    Không chứa logic xử lý, chỉ điều phối và chuyển đổi dữ liệu.
    """
    def load_document(self, file_path: str, document_id: int = None, extra_metadata: Dict[str, Any] | None = None) -> List[Document]:
        """
        Tải và xử lý file, trả về định dạng List[Document] mà RAGService cần.
        
        Args:
            file_path: Đường dẫn file
            document_id: ID tài liệu trong database (dùng cho lưu ảnh)
            extra_metadata: Metadata bổ sung
        """
        source = os.path.basename(file_path)
        _, extension = os.path.splitext(file_path.lower())
        
        processed_chunks: List[Dict] = []
        if extension == ".pdf":
            processed_chunks = process_pdf(file_path, document_id)
        elif extension == ".docx":
            processed_chunks = process_docx(file_path, document_id)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
            
        # Chuyển đổi từ List[Dict] sang List[Document]
        langchain_docs: List[Document] = []
        for i, chunk in enumerate(processed_chunks):
            metadata = {
                "source": source,
                "chunk": chunk.get("id", i),
                "page_start": chunk.get("page_start"),
                "page_end": chunk.get("page_end"),
                "has_ocr": chunk.get("has_ocr", False),
                "image_paths": chunk.get("image_paths", [])
            }
            if extra_metadata:
                metadata.update(extra_metadata)
            # Remove None values from metadata
            metadata = {k: v for k, v in metadata.items() if v is not None}
            
            doc = Document(page_content=chunk["text"], metadata=metadata)
            langchain_docs.append(doc)
            
        return langchain_docs
